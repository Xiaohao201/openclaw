from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "工作周报_2026-08-03至2026-08-07.md"
OUTPUT = ROOT / "浩苇工作周报_08-03至08-07.docx"

FONT = "Microsoft YaHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"
TABLE_WIDTHS = [3600, 4320, 1440]


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, 9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167
    number = doc.styles["List Number"]
    number.font.name = FONT
    number._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    number.font.size = Pt(11)
    number.paragraph_format.left_indent = Inches(0.5)
    number.paragraph_format.first_line_indent = Inches(-0.25)
    number.paragraph_format.space_after = Pt(8)
    number.paragraph_format.line_spacing = 1.167


def add_rich_text(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        set_run_font(run, 11, bold=bold)


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [c.strip() for c in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r"[-: ]+", c) for c in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("浩苇工作周报｜2026年8月3日—8月7日")
    set_run_font(run, 9, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    set_run_font(run, 20, bold=True, color="1F3A5F")

    i = 1
    while i < len(lines):
        raw = lines[i]
        text = raw.strip()
        if not text or text == "---":
            i += 1
            continue
        if text.startswith("## "):
            doc.add_paragraph(text[3:], style="Heading 1")
        elif text.startswith("### "):
            doc.add_paragraph(text[4:], style="Heading 2")
        elif text.startswith("|"):
            rows, i = parse_table(lines, i)
            table = doc.add_table(rows=len(rows), cols=3)
            table.style = "Table Grid"
            table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row[:3]):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.10
                    if c_idx == 2:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_rich_text(para, value)
                    for run in para.runs:
                        set_run_font(run, 9.5, bold=(r_idx == 0))
                    if r_idx == 0:
                        shade_cell(cell, LIGHT_GRAY)
            set_table_geometry(table, TABLE_WIDTHS)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        elif re.match(r"^\d+\. ", text):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\. ", "", text))
        elif text.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, text[2:])
        else:
            p = doc.add_paragraph()
            add_rich_text(p, text)
        i += 1

    doc.core_properties.title = "工作周报（2026年8月3日 – 8月7日）"
    doc.core_properties.subject = "上一周工作总结"
    doc.core_properties.author = "浩苇"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
